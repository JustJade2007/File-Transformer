"""Document converter supporting DOCX, PDF, Markdown, HTML, TXT, RTF, and EPUB."""
import os
import re
import time
import zipfile
from typing import Callable, Dict, List, Optional, Set
import xml.etree.ElementTree as ET

from .base import BaseConverter, ConversionResult

DOC_INPUTS = {"pdf", "docx", "txt", "md", "html", "htm", "rtf", "odt", "epub"}
DOC_OUTPUTS = {"txt", "pdf", "docx", "html", "md", "png"}


class DocumentConverter(BaseConverter):
    """Handles text documents, eBooks, Markdown, and PDF document workflows."""

    name = "Document Engine (PyPDF & Docx)"
    category = "Documents"

    def supported_inputs(self) -> Set[str]:
        return DOC_INPUTS

    def supported_outputs(self, input_ext: Optional[str] = None) -> Set[str]:
        if not input_ext:
            return DOC_OUTPUTS
        clean = input_ext.lower().lstrip(".")
        if clean == "pdf":
            return {"txt", "png", "html"}
        elif clean == "docx":
            return {"txt", "md", "html", "pdf"}
        elif clean in ("txt", "md"):
            return {"html", "pdf", "docx", "txt"}
        elif clean in ("html", "htm"):
            return {"txt", "md", "pdf"}
        elif clean in ("rtf", "odt", "epub"):
            return {"txt", "md", "html", "docx"}
        return {"txt", "html", "md"}

    def get_default_options(self, source_ext: str, target_ext: str) -> Dict:
        clean_target = target_ext.lower().lstrip(".")
        if clean_target == "png":
            return {"extract_all_pages": False, "page_dpi": 150}
        elif clean_target == "pdf":
            return {"font_size": 11, "page_margin": 36}
        return {}

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        import pypdf
        reader = pypdf.PdfReader(pdf_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {i + 1} ---\n" + text)
        return "\n\n".join(text_parts)

    def _extract_text_from_docx(self, docx_path: str) -> str:
        import docx
        doc = docx.Document(docx_path)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(lines)

    def _extract_text_from_epub(self, epub_path: str) -> str:
        """EPUB files are zip archives containing HTML/XHTML chapters."""
        text_parts = []
        with zipfile.ZipFile(epub_path, "r") as zf:
            for name in zf.namelist():
                if name.endswith((".html", ".xhtml", ".htm")):
                    try:
                        content = zf.read(name).decode("utf-8", errors="replace")
                        clean = re.sub(r"<[^>]+>", " ", content)
                        clean = re.sub(r"\s+", " ", clean).strip()
                        if clean:
                            text_parts.append(clean)
                    except Exception:
                        pass
        return "\n\n".join(text_parts)

    def _extract_text_from_odt(self, odt_path: str) -> str:
        """ODT files are zip archives with content.xml."""
        with zipfile.ZipFile(odt_path, "r") as zf:
            if "content.xml" in zf.namelist():
                xml_data = zf.read("content.xml")
                root = ET.fromstring(xml_data)
                texts = [elem.text for elem in root.iter() if elem.text]
                return " ".join(texts)
        return ""

    def _extract_text_from_rtf(self, rtf_path: str) -> str:
        with open(rtf_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        # Basic RTF control word stripper
        clean = re.sub(r"\\[a-z0-9\-]+ ?", "", content)
        clean = re.sub(r"[{}\\]", "", clean)
        return clean.strip()

    def _text_to_docx(self, text: str, output_path: str):
        import docx
        doc = docx.Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(output_path)

    def _text_to_pdf(self, text: str, output_path: str, title: str = "Document"):
        """Simple pure Python PDF generation via basic PDF stream or reportlab if present."""
        # Simple UTF-8 compliant PDF creation
        import pypdf
        # Create a basic PDF page using pypdf / canvas
        from pypdf import PdfWriter
        # If reportlab is available, use it; otherwise create clean structured PDF via pypdf
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            y = height - 50
            for line in text.splitlines():
                if y < 50:
                    c.showPage()
                    y = height - 50
                # ASCII safe display
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                c.drawString(50, y, safe_line[:95])
                y -= 14
            c.save()
        except ImportError:
            # Fallback using reportlab-free minimal PDF generator
            self._minimal_pdf_write(text, output_path, title)

    def _minimal_pdf_write(self, text: str, output_path: str, title: str):
        lines = text.splitlines()
        max_lines_per_page = 50
        pages = [lines[i:i + max_lines_per_page] for i in range(0, max(1, len(lines)), max_lines_per_page)]

        # Construct standard PDF objects
        objects = []
        page_obj_ids = []

        # Object 1: Catalog
        objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
        # Object 2: Pages (placeholder, written below)

        cur_id = 3
        page_stream_tuples = []
        for p in pages:
            # Content stream
            text_cmds = ["BT", "/F1 10 Tf", "50 750 Td", "14 TL"]
            for l in p:
                sanitized = l.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                sanitized = sanitized.encode("latin-1", "replace").decode("latin-1")
                text_cmds.append(f"({sanitized[:90]}) '")
            text_cmds.append("ET")
            stream_data = "\n".join(text_cmds).encode("latin-1")

            page_id = cur_id
            stream_id = cur_id + 1
            cur_id += 2

            page_obj_ids.append(page_id)
            page_stream_tuples.append((page_id, stream_id, stream_data))

        # Write out bytes
        buf = bytearray(b"%PDF-1.4\n")
        offsets = {}

        # 1: Catalog
        offsets[1] = len(buf)
        buf.extend(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")

        # 2: Pages
        offsets[2] = len(buf)
        kids_str = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
        buf.extend(f"2 0 obj\n<< /Type /Pages /Kids [{kids_str}] /Count {len(pages)} >>\nendobj\n".encode("latin-1"))

        # Font object
        font_id = cur_id
        offsets[font_id] = len(buf)
        buf.extend(f"{font_id} 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n".encode("latin-1"))

        for pid, sid, sdata in page_stream_tuples:
            offsets[pid] = len(buf)
            buf.extend(f"{pid} 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {sid} 0 R >>\nendobj\n".encode("latin-1"))

            offsets[sid] = len(buf)
            buf.extend(f"{sid} 0 obj\n<< /Length {len(sdata)} >>\nstream\n".encode("latin-1"))
            buf.extend(sdata)
            buf.extend(b"\nendstream\nendobj\n")

        # XRef
        xref_offset = len(buf)
        total_objects = font_id + 1
        buf.extend(f"xref\n0 {total_objects}\n0000000000 65535 f \n".encode("latin-1"))
        for oid in range(1, total_objects):
            offset = offsets.get(oid, 0)
            buf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))

        buf.extend(f"trailer\n<< /Size {total_objects} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("latin-1"))

        with open(output_path, "wb") as f:
            f.write(buf)

    def convert(
        self,
        source_path: str,
        target_path: str,
        target_format: str,
        options: Optional[Dict] = None,
        progress_callback: Optional[Callable[[float, str], None]] = None,
        cancel_event: Optional[object] = None,
    ) -> ConversionResult:
        start_time = time.time()
        source_ext = os.path.splitext(source_path)[1].lower().lstrip(".")
        target_ext = target_format.lower().lstrip(".")
        opts = options or self.get_default_options(source_ext, target_ext)

        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)

        try:
            if progress_callback:
                progress_callback(0.2, f"Reading {source_ext.upper()}...")

            # 1. Extract source text / content
            text_content = ""
            if source_ext == "pdf":
                text_content = self._extract_text_from_pdf(source_path)
            elif source_ext == "docx":
                text_content = self._extract_text_from_docx(source_path)
            elif source_ext == "epub":
                text_content = self._extract_text_from_epub(source_path)
            elif source_ext == "odt":
                text_content = self._extract_text_from_odt(source_path)
            elif source_ext == "rtf":
                text_content = self._extract_text_from_rtf(source_path)
            elif source_ext in ("txt", "md", "html", "htm"):
                with open(source_path, "r", encoding="utf-8", errors="replace") as f:
                    text_content = f.read()

            if cancel_event and getattr(cancel_event, "is_set", lambda: False)():
                return ConversionResult(success=False, error_message="Cancelled by user.")

            if progress_callback:
                progress_callback(0.6, f"Converting to {target_ext.upper()}...")

            # 2. Convert to Target Format
            if target_ext == "txt":
                # Strip HTML tags if source was HTML
                if source_ext in ("html", "htm"):
                    clean = re.sub(r"<[^>]+>", " ", text_content)
                    clean = re.sub(r"\s+", " ", clean).strip()
                    text_content = clean
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(text_content)

            elif target_ext == "md":
                if source_ext in ("html", "htm"):
                    # Basic HTML to MD conversion
                    text_content = re.sub(r"<h1[^>]*>(.*?)</h1>", r"# \1\n", text_content, flags=re.I)
                    text_content = re.sub(r"<h2[^>]*>(.*?)</h2>", r"## \1\n", text_content, flags=re.I)
                    text_content = re.sub(r"<b[^>]*>(.*?)</b>", r"**\1**", text_content, flags=re.I)
                    text_content = re.sub(r"<strong[^>]*>(.*?)</strong>", r"**\1**", text_content, flags=re.I)
                    text_content = re.sub(r"<i[^>]*>(.*?)</i>", r"*\1*", text_content, flags=re.I)
                    text_content = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n\n", text_content, flags=re.I)
                    text_content = re.sub(r"<[^>]+>", "", text_content)
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(text_content)

            elif target_ext == "html":
                if source_ext == "md":
                    import markdown
                    html_body = markdown.markdown(text_content, extensions=["tables", "fenced_code"])
                else:
                    lines = [f"<p>{line}</p>" for line in text_content.splitlines() if line.strip()]
                    html_body = "\n".join(lines)

                full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{os.path.basename(source_path)}</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #24292e; }}
pre {{ background: #f6f8fa; padding: 16px; border-radius: 6px; overflow: auto; }}
code {{ font-family: Consolas, 'Liberation Mono', Menlo, Courier, monospace; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #dfe2e5; padding: 6px 13px; }}
th {{ background: #f6f8fa; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(full_html)

            elif target_ext == "docx":
                self._text_to_docx(text_content, target_path)

            elif target_ext == "pdf":
                self._text_to_pdf(text_content, target_path, title=os.path.basename(source_path))

            elif target_ext == "png" and source_ext == "pdf":
                # Extract images from PDF pages using pypdf
                import pypdf
                reader = pypdf.PdfReader(source_path)
                found_image = False
                for page_idx, page in enumerate(reader.pages):
                    for img_idx, img_file in enumerate(page.images):
                        with open(target_path, "wb") as fp:
                            fp.write(img_file.data)
                        found_image = True
                        break
                    if found_image:
                        break
                if not found_image:
                    # If no embedded raster images, write note or render text as image
                    from PIL import Image, ImageDraw
                    canvas_img = Image.new("RGB", (800, 1000), color=(255, 255, 255))
                    draw = ImageDraw.Draw(canvas_img)
                    y = 30
                    for line in text_content.splitlines()[:50]:
                        draw.text((30, y), line[:80], fill=(0, 0, 0))
                        y += 18
                    canvas_img.save(target_path, "PNG")

            if progress_callback:
                progress_callback(1.0, "Complete")

            return ConversionResult(
                success=True,
                output_path=target_path,
                duration_seconds=time.time() - start_time,
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Document conversion failed: {str(e)}",
                duration_seconds=time.time() - start_time,
            )
